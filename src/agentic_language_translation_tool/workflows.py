"""End-to-end workspace workflows for translations, QA, and rebuilds."""

from __future__ import annotations

import json
import re
import shutil
from pathlib import Path
from typing import TypeGuard, cast

from agentic_language_translation_tool.adversarial_qa import detect_adversarial_qa_issues
from agentic_language_translation_tool.glossary import (
    GLOSSARY_FILE,
    enrich_segments_with_glossary,
    find_glossary_entries,
    likely_terminology_drift,
    load_glossary,
    read_workspace_glossary,
    term_in_text,
)
from agentic_language_translation_tool.io import (
    atomic_write_json,
    atomic_write_text,
    file_sha256,
    read_jsonl_model,
    stable_id,
    write_jsonl,
)
from agentic_language_translation_tool.models import (
    Batch,
    BatchPurpose,
    BatchStatus,
    FindingSeverity,
    Glossary,
    GlossaryEntry,
    GlossaryMetadata,
    JobManifest,
    JobStage,
    JobState,
    QaIssue,
    QaReport,
    Segment,
    TranslationRecord,
    VerificationFinding,
)
from agentic_language_translation_tool.tasks import (
    render_correction_batch,
    render_translation_batch,
    render_verification_batch,
)
from agentic_language_translation_tool.workspace import (
    SEGMENTS_FILE,
    STRUCTURE_FILE,
    WorkspaceError,
    load_manifest,
    read_segments,
    save_manifest,
    write_resume,
)

QA_REPORT_JSON = "qa_report.json"
QA_REPORT_MD = "qa_report.md"
MARKDOWN_LINK_PATTERN = re.compile(r"\[[^\]]+\]\(([^)]+)\)")
OBJECTIVE_DUPLICATE_SUPPRESSORS = {
    "placeholder_drift",
    "protected_term_damage",
    "required_translation_missing",
    "preferred_translation_missing",
    "forbidden_translation_used",
    "do_not_translate_term_changed",
    "markdown_link_damage",
}


def apply_glossary(workspace: Path, glossary: Path) -> JobManifest:
    """Apply or refresh a workspace glossary and regenerate task files."""
    manifest = load_manifest(workspace)
    normalized_glossary = load_glossary(glossary)
    segments = enrich_segments_with_glossary(read_segments(workspace), normalized_glossary)
    atomic_write_json(workspace / GLOSSARY_FILE, normalized_glossary)
    write_jsonl(workspace / SEGMENTS_FILE, segments)
    manifest.glossary = GlossaryMetadata(
        source_file=normalized_glossary.source_file or glossary.name,
        source_checksum=normalized_glossary.source_checksum or file_sha256(glossary),
        applied_at=normalized_glossary.applied_at,
    )
    _regenerate_batch_files(workspace, manifest, segments)
    _save_and_resume(workspace, manifest)
    return manifest


def apply_translations(workspace: Path, translations: Path, *, force: bool = False) -> JobManifest:
    """Apply agent-produced translation records to workspace segments."""
    manifest = load_manifest(workspace)
    segments = read_segments(workspace)
    records = read_jsonl_model(translations, TranslationRecord)
    if not records:
        raise WorkspaceError(f"translation file has no records: {translations}")

    segments_by_id = {segment.segment_id: segment for segment in segments}
    seen: set[str] = set()
    for record in records:
        if record.segment_id not in segments_by_id:
            raise WorkspaceError(f"unknown segment_id in translations: {record.segment_id}")
        if record.segment_id in seen:
            raise WorkspaceError(f"duplicate segment_id in translations: {record.segment_id}")
        segment = segments_by_id[record.segment_id]
        if segment.translated_text and not force:
            raise WorkspaceError(f"segment already has translation: {record.segment_id}")
        segment.translated_text = record.translated_text
        if record.notes:
            segment.notes = sorted(set(segment.notes + record.notes))
        seen.add(record.segment_id)

    _persist_translation_input(workspace, translations, records)
    write_jsonl(workspace / SEGMENTS_FILE, segments)
    _refresh_translation_batch_statuses(manifest, segments)
    verification_planned = any(
        batch.purpose == BatchPurpose.VERIFICATION for batch in manifest.batches
    )
    manifest.state.stage = (
        JobStage.VERIFICATION_PLANNED
        if _all_segments_translated(segments) and verification_planned
        else JobStage.PARTIALLY_TRANSLATED
    )
    _save_and_resume(workspace, manifest)
    return manifest


def apply_verification(workspace: Path, results: Path) -> JobManifest:
    """Apply verifier findings to a workspace and update verification batch states."""
    manifest = load_manifest(workspace)
    segments = read_segments(workspace)
    findings = read_jsonl_model(results, VerificationFinding)
    segment_ids = {segment.segment_id for segment in segments}
    for finding in findings:
        if finding.segment_id not in segment_ids:
            raise WorkspaceError(
                f"unknown segment_id in verification results: {finding.segment_id}"
            )

    _persist_verification_input(workspace, results, findings)
    existing_by_id = {finding.finding_id: finding for finding in manifest.findings}
    for finding in findings:
        existing_by_id[finding.finding_id] = finding
    manifest.findings = list(existing_by_id.values())
    _refresh_verification_batch_statuses(manifest)
    if any(finding.severity == FindingSeverity.BLOCKER for finding in manifest.findings):
        manifest.state.stage = JobStage.BLOCKED
    elif _has_unresolved_correction_findings(manifest):
        manifest.state.stage = JobStage.CORRECTIONS_PLANNED
    else:
        manifest.state.stage = JobStage.VERIFICATION_PLANNED
    _save_and_resume(workspace, manifest)
    return manifest


def plan_corrections(workspace: Path) -> JobManifest:
    """Create correction batches for unresolved warning/error/blocker findings."""
    manifest = load_manifest(workspace)
    segments = read_segments(workspace)
    findings_by_segment = _actionable_findings_by_segment(manifest.findings)
    if not findings_by_segment:
        _save_and_resume(workspace, manifest)
        return manifest

    segments_by_id = {segment.segment_id: segment for segment in segments}
    existing_correction_ids = {
        batch.batch_id for batch in manifest.batches if batch.purpose == BatchPurpose.CORRECTION
    }
    correction_batches: list[Batch] = []
    for index, segment_id in enumerate(sorted(findings_by_segment), start=1):
        batch_id = f"correction_{index:04d}_{stable_id(segment_id, length=8)}"
        if batch_id in existing_correction_ids:
            continue
        segment = segments_by_id[segment_id]
        file_name = f"{batch_id}.md"
        batch = Batch(
            batch_id=batch_id,
            purpose=BatchPurpose.CORRECTION,
            segment_ids=[segment_id],
            token_estimate=max(1, len(segment.source_text + (segment.translated_text or "")) // 4),
            status=BatchStatus.PENDING,
            file=f"correction_batches/{file_name}",
        )
        content = render_correction_batch(batch, segment, findings_by_segment[segment_id])
        atomic_write_text(workspace / batch.file, content)
        correction_batches.append(batch)
    manifest.batches.extend(correction_batches)
    manifest.state.stage = JobStage.CORRECTIONS_PLANNED
    _save_and_resume(workspace, manifest)
    return manifest


def validate_job(workspace: Path) -> QaReport:
    """Run deterministic QA checks and write QA reports."""
    manifest = load_manifest(workspace)
    segments = read_segments(workspace)
    glossary = read_workspace_glossary(workspace)
    issues: list[QaIssue] = []
    issues.extend(
        _translation_issues(
            segments,
            glossary,
            source_language=manifest.source_language,
            target_language=manifest.target_language,
        )
    )
    issues.extend(_verification_issues(manifest, segments))
    issues.extend(_finding_issues(manifest.findings))

    passed = not any(
        issue.severity in {FindingSeverity.ERROR, FindingSeverity.BLOCKER}
        for issue in issues
    )
    report = QaReport(
        job_id=manifest.job_id,
        workspace_path=str(workspace.resolve()),
        passed=passed,
        issue_count=len(issues),
        issues=issues,
    )
    atomic_write_json(workspace / QA_REPORT_JSON, report)
    atomic_write_text(workspace / QA_REPORT_MD, _render_qa_report(report))
    manifest.state.stage = JobStage.VALIDATED if passed else JobStage.BLOCKED
    _save_and_resume(workspace, manifest)
    return report


def rebuild_document(workspace: Path, output: Path, *, force: bool = False) -> Path:
    """Rebuild TXT or Markdown output from translated segments."""
    manifest = load_manifest(workspace)
    segments = read_segments(workspace)
    blockers = [
        finding for finding in manifest.findings if finding.severity == FindingSeverity.BLOCKER
    ]
    missing = [segment.segment_id for segment in segments if not segment.translated_text]
    if (missing or blockers) and not force:
        reasons = []
        if missing:
            reasons.append(f"missing translations: {', '.join(missing)}")
        if blockers:
            reasons.append("unresolved blocker findings")
        raise WorkspaceError("; ".join(reasons))

    structure = _read_structure(workspace)
    document_format = str(structure.get("format", ""))
    if document_format == "txt":
        content = _rebuild_txt(segments)
        atomic_write_text(output, content)
    elif document_format == "markdown":
        content = _rebuild_markdown(structure, segments)
        atomic_write_text(output, content)
    elif document_format == "docx":
        _rebuild_docx(workspace, structure, segments, output)
    elif document_format == "pdf":
        if output.suffix.lower() == ".pdf" and not force:
            raise WorkspaceError(
                "exact visual PDF recreation is future work; choose a Markdown output path "
                "or pass --force to write translated Markdown content anyway"
            )
        content = _rebuild_pdf_markdown(structure, segments)
        atomic_write_text(output, content)
    else:
        raise WorkspaceError(f"rebuild unsupported for format: {document_format}")
    manifest.state.stage = JobStage.REBUILT
    _save_and_resume(workspace, manifest)
    return output


def _persist_translation_input(
    workspace: Path,
    translations: Path,
    records: list[TranslationRecord],
) -> None:
    target = workspace / "translations" / f"{translations.stem}.jsonl"
    if translations.resolve() == target.resolve():
        write_jsonl(target, records)
    else:
        write_jsonl(target, records)


def _persist_verification_input(
    workspace: Path,
    results: Path,
    findings: list[VerificationFinding],
) -> None:
    target = workspace / "verification_results" / f"{results.stem}.jsonl"
    if results.resolve() == target.resolve():
        write_jsonl(target, findings)
    else:
        write_jsonl(target, findings)


def _regenerate_batch_files(
    workspace: Path,
    manifest: JobManifest,
    segments: list[Segment],
) -> None:
    segments_by_id = {segment.segment_id: segment for segment in segments}
    findings_by_segment = _actionable_findings_by_segment(manifest.findings)
    for batch in manifest.batches:
        batch_segments = [segments_by_id[segment_id] for segment_id in batch.segment_ids]
        path = workspace / batch.file
        if batch.purpose == BatchPurpose.TRANSLATION:
            content = render_translation_batch(batch, batch_segments)
        elif batch.purpose == BatchPurpose.VERIFICATION:
            content = render_verification_batch(batch, batch_segments)
        elif batch.purpose == BatchPurpose.CORRECTION and len(batch_segments) == 1:
            segment = batch_segments[0]
            content = render_correction_batch(
                batch,
                segment,
                findings_by_segment.get(segment.segment_id, []),
            )
        else:
            continue
        atomic_write_text(path, content)


def _refresh_translation_batch_statuses(manifest: JobManifest, segments: list[Segment]) -> None:
    translated_ids = {segment.segment_id for segment in segments if segment.translated_text}
    for batch in manifest.batches:
        if batch.purpose == BatchPurpose.TRANSLATION and set(batch.segment_ids) <= translated_ids:
            batch.status = BatchStatus.TRANSLATED


def _refresh_verification_batch_statuses(manifest: JobManifest) -> None:
    severity_by_segment: dict[str, FindingSeverity] = {}
    for finding in manifest.findings:
        current = severity_by_segment.get(finding.segment_id)
        if current is None or _severity_rank(finding.severity) > _severity_rank(current):
            severity_by_segment[finding.segment_id] = finding.severity

    for batch in manifest.batches:
        if batch.purpose != BatchPurpose.VERIFICATION:
            continue
        severities = [
            severity_by_segment[segment_id]
            for segment_id in batch.segment_ids
            if segment_id in severity_by_segment
        ]
        if any(severity == FindingSeverity.BLOCKER for severity in severities):
            batch.status = BatchStatus.BLOCKED
        elif any(
            severity in {FindingSeverity.WARNING, FindingSeverity.ERROR}
            for severity in severities
        ):
            batch.status = BatchStatus.NEEDS_CORRECTION
        else:
            batch.status = BatchStatus.VERIFIED


def _all_segments_translated(segments: list[Segment]) -> bool:
    return all(segment.translated_text for segment in segments)


def _has_unresolved_correction_findings(manifest: JobManifest) -> bool:
    actionable = {FindingSeverity.WARNING, FindingSeverity.ERROR, FindingSeverity.BLOCKER}
    return any(
        finding.severity in actionable
        for finding in manifest.findings
    )


def _actionable_findings_by_segment(
    findings: list[VerificationFinding],
) -> dict[str, list[VerificationFinding]]:
    grouped: dict[str, list[VerificationFinding]] = {}
    actionable = {FindingSeverity.WARNING, FindingSeverity.ERROR, FindingSeverity.BLOCKER}
    for finding in findings:
        if finding.severity in actionable:
            grouped.setdefault(finding.segment_id, []).append(finding)
    return grouped


def _translation_issues(
    segments: list[Segment],
    glossary: Glossary | None,
    *,
    source_language: str = "",
    target_language: str = "",
) -> list[QaIssue]:
    issues: list[QaIssue] = []
    for segment in segments:
        segment_issues: list[QaIssue] = []
        translation = segment.translated_text or ""
        if not translation.strip():
            segment_issues.append(
                _issue(
                    segment.segment_id,
                    "missing_translation",
                    FindingSeverity.ERROR,
                    "Missing translation.",
                )
            )
            issues.extend(segment_issues)
            continue
        source_placeholders = set(segment.placeholders)
        missing_placeholders = sorted(
            token for token in source_placeholders if token not in translation
        )
        if missing_placeholders:
            segment_issues.append(
                _issue(
                    segment.segment_id,
                    "placeholder_drift",
                    FindingSeverity.ERROR,
                    f"Missing placeholders: {', '.join(missing_placeholders)}.",
                )
            )
        missing_terms = sorted(term for term in segment.protected_terms if term not in translation)
        if missing_terms:
            segment_issues.append(
                _issue(
                    segment.segment_id,
                    "protected_term_damage",
                    FindingSeverity.ERROR,
                    f"Missing protected terms: {', '.join(missing_terms)}.",
                )
            )
        if segment.source_text.strip() == translation.strip():
            segment_issues.append(
                _issue(
                    segment.segment_id,
                    "untranslated_text",
                    FindingSeverity.WARNING,
                    "Translation is identical to the source text.",
                )
            )
        if _has_length_anomaly(segment.source_text, translation):
            segment_issues.append(
                _issue(
                    segment.segment_id,
                    "length_anomaly",
                    FindingSeverity.WARNING,
                    "Translation length differs sharply from the source segment.",
                )
            )
        if "markdown" in segment.format:
            segment_issues.extend(_markdown_link_issues(segment, translation))
        if glossary is not None:
            segment_issues.extend(_terminology_issues(segment, glossary, translation))
        segment_issues.extend(
            _filtered_adversarial_issues(
                segment,
                translation,
                segment_issues,
                source_language=source_language,
                target_language=target_language,
            )
        )
        issues.extend(segment_issues)
    return issues


def _filtered_adversarial_issues(
    segment: Segment,
    translation: str,
    existing_issues: list[QaIssue],
    *,
    source_language: str,
    target_language: str,
) -> list[QaIssue]:
    existing_categories = {issue.category for issue in existing_issues}
    adversarial_issues = detect_adversarial_qa_issues(
        segment,
        translation,
        source_language=source_language,
        target_language=target_language,
    )
    if existing_categories.intersection(OBJECTIVE_DUPLICATE_SUPPRESSORS):
        return [
            issue
            for issue in adversarial_issues
            if issue.category
            not in {"possible_entity_damage", "possible_omission", "possible_overly_literal"}
        ]
    return adversarial_issues


def _terminology_issues(
    segment: Segment,
    glossary: Glossary,
    translation: str,
) -> list[QaIssue]:
    issues: list[QaIssue] = []
    for entry in find_glossary_entries(segment.source_text, glossary):
        if entry.do_not_translate and not term_in_text(entry.source_term, translation, entry):
            issues.append(
                _issue(
                    segment.segment_id,
                    "do_not_translate_term_changed",
                    FindingSeverity.ERROR,
                    f"Do-not-translate term changed or missing: {entry.source_term}.",
                )
            )
        if entry.required_translation and not _contains_rule_text(
            translation,
            entry.required_translation,
            entry,
        ):
            issues.append(
                _issue(
                    segment.segment_id,
                    "required_translation_missing",
                    FindingSeverity.ERROR,
                    f"Missing required translation: {entry.required_translation}.",
                )
            )
            issues.extend(_drift_issue(segment, entry, entry.required_translation, translation))
        if (
            entry.preferred_translation
            and entry.preferred_translation != entry.required_translation
            and not _contains_rule_text(translation, entry.preferred_translation, entry)
        ):
            issues.append(
                _issue(
                    segment.segment_id,
                    "preferred_translation_missing",
                    FindingSeverity.WARNING,
                    f"Missing preferred translation: {entry.preferred_translation}.",
                )
            )
            issues.extend(_drift_issue(segment, entry, entry.preferred_translation, translation))
        for forbidden in entry.forbidden_translations:
            if _contains_rule_text(translation, forbidden, entry):
                issues.append(
                    _issue(
                        segment.segment_id,
                        "forbidden_translation_used",
                        FindingSeverity.ERROR,
                        f"Forbidden translation used: {forbidden}.",
                    )
                )
    return issues


def _drift_issue(
    segment: Segment,
    entry: GlossaryEntry,
    expected: str,
    translation: str,
) -> list[QaIssue]:
    if not likely_terminology_drift(expected, translation):
        return []
    return [
        _issue(
            segment.segment_id,
            "terminology_drift",
            FindingSeverity.WARNING,
            f"Possible terminology drift for {entry.source_term}; expected {expected}.",
        )
    ]


def _contains_rule_text(text: str, expected: str, entry: GlossaryEntry) -> bool:
    if entry.case_sensitive:
        return expected in text
    return expected.casefold() in text.casefold()


def _verification_issues(manifest: JobManifest, segments: list[Segment]) -> list[QaIssue]:
    verified_segment_ids: set[str] = set()
    for batch in manifest.batches:
        if batch.purpose != BatchPurpose.VERIFICATION:
            continue
        verified_statuses = {
            BatchStatus.VERIFIED,
            BatchStatus.NEEDS_CORRECTION,
            BatchStatus.BLOCKED,
        }
        if batch.status in verified_statuses:
            verified_segment_ids.update(batch.segment_ids)

    return [
        _issue(
            segment.segment_id,
            "missing_verification",
            FindingSeverity.WARNING,
            "Translated segment has not been marked verified by a verifier batch.",
        )
        for segment in segments
        if segment.translated_text and segment.segment_id not in verified_segment_ids
    ]


def _finding_issues(findings: list[VerificationFinding]) -> list[QaIssue]:
    actionable = {FindingSeverity.WARNING, FindingSeverity.ERROR, FindingSeverity.BLOCKER}
    return [
        _issue(
            finding.segment_id,
            f"verification_{finding.category.value}",
            finding.severity,
            finding.explanation,
        )
        for finding in findings
        if finding.severity in actionable
    ]


def _markdown_link_issues(segment: Segment, translation: str) -> list[QaIssue]:
    source_urls = set(MARKDOWN_LINK_PATTERN.findall(segment.source_text))
    translated_urls = set(MARKDOWN_LINK_PATTERN.findall(translation))
    missing_urls = sorted(source_urls - translated_urls)
    if not missing_urls:
        return []
    return [
        _issue(
            segment.segment_id,
            "markdown_link_damage",
            FindingSeverity.ERROR,
            f"Missing Markdown link targets: {', '.join(missing_urls)}.",
        )
    ]


def _has_length_anomaly(source: str, translation: str) -> bool:
    source_len = max(1, len(source.strip()))
    translation_len = len(translation.strip())
    ratio = translation_len / source_len
    return ratio < 0.25 or ratio > 2.5


def _issue(
    segment_id: str | None,
    category: str,
    severity: FindingSeverity,
    explanation: str,
) -> QaIssue:
    issue_id = f"qa_{stable_id(str(segment_id), category, explanation)}"
    return QaIssue(
        issue_id=issue_id,
        segment_id=segment_id,
        category=category,
        severity=severity,
        explanation=explanation,
    )


def _render_qa_report(report: QaReport) -> str:
    lines = [
        "# QA Report",
        "",
        f"- Job ID: `{report.job_id}`",
        f"- Passed: `{str(report.passed).lower()}`",
        f"- Issue count: `{report.issue_count}`",
        "",
        "## Issues",
        "",
    ]
    if not report.issues:
        lines.append("- None.")
    else:
        for issue in report.issues:
            segment = issue.segment_id or "job"
            lines.append(
                f"- `{issue.severity.value}` `{issue.category}` `{segment}`: {issue.explanation}"
            )
    return "\n".join(lines).rstrip() + "\n"


def _severity_rank(severity: FindingSeverity) -> int:
    ranks = {
        FindingSeverity.INFO: 0,
        FindingSeverity.WARNING: 1,
        FindingSeverity.ERROR: 2,
        FindingSeverity.BLOCKER: 3,
    }
    return ranks[severity]


def _read_structure(workspace: Path) -> dict[str, object]:
    raw_structure = json.loads((workspace / STRUCTURE_FILE).read_text(encoding="utf-8"))
    return cast(dict[str, object], raw_structure)


def _rebuild_txt(segments: list[Segment]) -> str:
    return (
        "\n\n".join(segment.translated_text or segment.source_text for segment in segments)
        + "\n"
    )


def _rebuild_markdown(structure: dict[str, object], segments: list[Segment]) -> str:
    blocks = structure.get("blocks")
    if not isinstance(blocks, list):
        raise WorkspaceError("invalid markdown structure: missing blocks")
    translated_by_id = {segment.segment_id: segment.translated_text for segment in segments}
    rebuilt: list[str] = []
    for block in blocks:
        if not isinstance(block, dict):
            raise WorkspaceError("invalid markdown structure block")
        text = block.get("text")
        segment_id = block.get("segment_id")
        translatable = block.get("translatable")
        if not isinstance(text, str) or not isinstance(translatable, bool):
            raise WorkspaceError("invalid markdown structure block fields")
        if translatable:
            if not isinstance(segment_id, str):
                raise WorkspaceError("translatable markdown block has no segment_id")
            rebuilt.append(translated_by_id.get(segment_id) or text)
        else:
            rebuilt.append(text)
    return "\n\n".join(rebuilt).rstrip() + "\n"


def _rebuild_docx(
    workspace: Path,
    structure: dict[str, object],
    segments: list[Segment],
    output: Path,
) -> None:
    from docx import Document

    source_file = structure.get("source_file")
    if not isinstance(source_file, str):
        raise WorkspaceError("invalid docx structure: missing source_file")
    source_path = workspace / "source" / source_file
    if not source_path.exists():
        raise WorkspaceError(f"missing source DOCX in workspace: {source_path}")
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source_path, output)
    document = Document(str(output))
    translated_by_id = {segment.segment_id: segment.translated_text for segment in segments}
    blocks = structure.get("blocks")
    if not isinstance(blocks, list):
        raise WorkspaceError("invalid docx structure: missing blocks")
    for block in blocks:
        if not isinstance(block, dict):
            raise WorkspaceError("invalid docx structure block")
        segment_id = block.get("segment_id")
        path_value = block.get("path")
        if not isinstance(segment_id, str) or not _is_path_list(path_value):
            raise WorkspaceError("invalid docx structure block fields")
        translation = translated_by_id.get(segment_id)
        if not translation:
            continue
        if path_value[0] == "paragraph":
            paragraph_index = int(path_value[1]) - 1
            _replace_paragraph_text(document.paragraphs[paragraph_index], translation)
        elif path_value[0] == "table":
            table_index = int(path_value[1]) - 1
            row_index = int(path_value[3]) - 1
            cell_index = int(path_value[5]) - 1
            cell = document.tables[table_index].rows[row_index].cells[cell_index]
            _replace_cell_text(cell, translation)
    document.save(str(output))


def _rebuild_pdf_markdown(structure: dict[str, object], segments: list[Segment]) -> str:
    translated_by_id = {segment.segment_id: segment.translated_text for segment in segments}
    pages = structure.get("pages")
    if not isinstance(pages, list):
        raise WorkspaceError("invalid pdf structure: missing pages")
    lines = ["# Translated PDF Text", ""]
    warnings = structure.get("warnings", [])
    if isinstance(warnings, list) and warnings:
        lines.extend(["## Extraction Warnings", ""])
        for warning in warnings:
            lines.append(f"- {warning}")
        lines.append("")
    for page in pages:
        if not isinstance(page, dict):
            raise WorkspaceError("invalid pdf page structure")
        page_number = page.get("page_number")
        blocks = page.get("blocks")
        if not isinstance(page_number, int) or not isinstance(blocks, list):
            raise WorkspaceError("invalid pdf page fields")
        lines.extend([f"## Page {page_number}", ""])
        for block in blocks:
            if not isinstance(block, dict):
                raise WorkspaceError("invalid pdf block structure")
            segment_id = block.get("segment_id")
            text = block.get("text")
            if not isinstance(segment_id, str) or not isinstance(text, str):
                raise WorkspaceError("invalid pdf block fields")
            lines.extend([translated_by_id.get(segment_id) or text, ""])
    return "\n".join(lines).rstrip() + "\n"


def _replace_paragraph_text(paragraph: object, text: str) -> None:
    runs = getattr(paragraph, "runs", [])
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)  # type: ignore[attr-defined]


def _replace_cell_text(cell: object, text: str) -> None:
    paragraphs = getattr(cell, "paragraphs", [])
    if not paragraphs:
        return
    _replace_paragraph_text(paragraphs[0], text)
    for paragraph in paragraphs[1:]:
        _replace_paragraph_text(paragraph, "")


def _is_path_list(value: object) -> TypeGuard[list[str]]:
    return isinstance(value, list) and all(isinstance(part, str) for part in value)


def _save_and_resume(workspace: Path, manifest: JobManifest) -> None:
    manifest.state = _state_for_manifest(manifest, workspace)
    save_manifest(workspace, manifest)
    write_resume(workspace, manifest)


def _state_for_manifest(manifest: JobManifest, workspace: Path) -> JobState:
    from agentic_language_translation_tool.workspace import _state_from_batches

    return _state_from_batches(manifest.state.stage, manifest.batches, workspace)


def copy_output_to_workspace(output: Path, workspace: Path) -> None:
    """Copy a rebuilt output into the workspace output folder."""
    target = workspace / "output" / output.name
    if output.resolve() != target.resolve():
        shutil.copy2(output, target)
