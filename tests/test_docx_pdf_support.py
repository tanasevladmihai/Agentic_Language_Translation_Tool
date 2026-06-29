from __future__ import annotations

from pathlib import Path

import fitz
import pytest
from docx import Document

from agentic_language_translation_tool.extractors import extract_document, inspect_input
from agentic_language_translation_tool.io import write_jsonl
from agentic_language_translation_tool.models import TranslationRecord
from agentic_language_translation_tool.workflows import apply_translations, rebuild_document
from agentic_language_translation_tool.workspace import (
    WorkspaceError,
    init_workspace,
    read_segments,
)


def test_docx_extraction_creates_segments_and_structural_paths(tmp_path: Path) -> None:
    source = _make_docx_fixture(tmp_path / "source.docx")

    result = extract_document(source)

    assert result.document_format == "docx"
    assert [segment.source_text for segment in result.segments] == [
        "Quarterly Update",
        "Hello team.",
        "Name",
        "Role",
        "Ada",
        "Engineer",
    ]
    assert result.segments[0].path == ["paragraph", "1"]
    assert result.segments[2].path == ["table", "1", "row", "1", "cell", "1"]
    assert "heading" in result.segments[0].style_tags


def test_docx_workspace_rebuild_preserves_table_structure(tmp_path: Path) -> None:
    source = _make_docx_fixture(tmp_path / "source.docx")
    workspace = tmp_path / "workspace"
    init_workspace(source, workspace, source_language="en", target_language="de")
    segments = read_segments(workspace)
    translations = tmp_path / "translations.jsonl"
    translated_text = {
        "Quarterly Update": "Quartalsbericht",
        "Hello team.": "Hallo Team.",
        "Name": "Name",
        "Role": "Rolle",
        "Ada": "Ada",
        "Engineer": "Ingenieurin",
    }
    write_jsonl(
        translations,
        [
            TranslationRecord(
                segment_id=segment.segment_id,
                translated_text=translated_text[segment.source_text],
            )
            for segment in segments
        ],
    )
    apply_translations(workspace, translations)
    output = tmp_path / "translated.docx"

    rebuild_document(workspace, output)

    rebuilt = Document(output)
    assert rebuilt.paragraphs[0].text == "Quartalsbericht"
    assert rebuilt.paragraphs[1].text == "Hallo Team."
    assert rebuilt.tables[0].rows[0].cells[1].text == "Rolle"
    assert rebuilt.tables[0].rows[1].cells[1].text == "Ingenieurin"


def test_pdf_extraction_creates_page_block_segments(tmp_path: Path) -> None:
    source = _make_pdf_fixture(tmp_path / "source.pdf")

    result = extract_document(source)

    assert result.document_format == "pdf"
    assert len(result.segments) == 2
    assert result.segments[0].path == ["page", "1", "block", "1"]
    assert result.segments[1].context == "source.pdf block 2"
    assert result.structure["page_count"] == 2


def test_blank_pdf_records_extraction_warning(tmp_path: Path) -> None:
    source = tmp_path / "blank.pdf"
    document = fitz.open()
    document.new_page()
    document.save(source)
    document.close()

    result = extract_document(source)

    assert result.segments == []
    warnings = result.structure["warnings"]
    assert isinstance(warnings, list)
    assert "page 1 has no extractable text blocks" in warnings


def test_pdf_workspace_rebuilds_translated_markdown_and_refuses_pdf(tmp_path: Path) -> None:
    source = _make_pdf_fixture(tmp_path / "source.pdf")
    workspace = tmp_path / "workspace"
    init_workspace(source, workspace, source_language="en", target_language="de")
    segments = read_segments(workspace)
    translations = tmp_path / "translations.jsonl"
    write_jsonl(
        translations,
        [
            TranslationRecord(
                segment_id=segment.segment_id,
                translated_text=f"DE: {segment.source_text}",
            )
            for segment in segments
        ],
    )
    apply_translations(workspace, translations)

    with pytest.raises(WorkspaceError, match="exact visual PDF recreation"):
        rebuild_document(workspace, tmp_path / "translated.pdf")

    output = tmp_path / "translated.md"
    rebuild_document(workspace, output)

    rebuilt = output.read_text(encoding="utf-8")
    assert "# Translated PDF Text" in rebuilt
    assert "## Page 1" in rebuilt
    assert "DE: First page text." in rebuilt
    assert "## Page 2" in rebuilt


def test_inspect_reports_docx_and_pdf_rebuild_capabilities(tmp_path: Path) -> None:
    docx_source = _make_docx_fixture(tmp_path / "source.docx")
    pdf_source = _make_pdf_fixture(tmp_path / "source.pdf")

    docx_metadata = inspect_input(docx_source)
    pdf_metadata = inspect_input(pdf_source)

    assert docx_metadata["same_format_rebuild_supported"] is True
    assert pdf_metadata["same_format_rebuild_supported"] is False
    assert "Text-layer extraction only; OCR is not implemented." in pdf_metadata["limitations"]


def _make_docx_fixture(path: Path) -> Path:
    document = Document()
    document.add_heading("Quarterly Update", level=1)
    document.add_paragraph("Hello team.")
    document.add_paragraph("")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Name"
    table.rows[0].cells[1].text = "Role"
    table.rows[1].cells[0].text = "Ada"
    table.rows[1].cells[1].text = "Engineer"
    document.save(path)
    return path


def _make_pdf_fixture(path: Path) -> Path:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "First page text.")
    second_page = document.new_page()
    second_page.insert_text((72, 72), "Second page text.")
    document.save(path)
    document.close()
    return path
